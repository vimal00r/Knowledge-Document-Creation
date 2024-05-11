import streamlit as st
import pandas as pd
import csv
import comtypes
comtypes.CoInitialize()
import os
import re
import PyPDF2
import base64

from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader
from nltk.tokenize import sent_tokenize
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from Paragraph_extraction_from_pdf import *
from Table_to_text.Detect_table_redact import *
from Flowchart_to_text.Detect_flowchart_redact import *
from Flowchart_to_text.flowchart_to_text import *



st.title("Knowledge Document Creation")

# For Main PDF
uploaded_main_pdf_document = st.file_uploader("Upload main PDF file(Ensure that the file is CITED)", type=["pdf"],)
main_pdf_folder_name = "main_pdf_files"
os.makedirs(main_pdf_folder_name, exist_ok=True)
if uploaded_main_pdf_document:
    main_document_path = ""
    with open(os.path.join(main_pdf_folder_name, uploaded_main_pdf_document.name), "wb") as f:
        f.write(uploaded_main_pdf_document.read())
        main_document_path = os.path.join(main_pdf_folder_name,uploaded_main_pdf_document.name)







with st.sidebar.title("PDF Viewer"):
    st.title("Pdf")
    if uploaded_main_pdf_document:
        def show_pdf(file_path):
            with open(file_path,"rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode('utf-8')
            pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="400" height="500" type="application/pdf">'
            st.markdown(pdf_display, unsafe_allow_html=True)
        show_pdf(main_document_path)


# For Reference PDF
def save_uploaded_pdf(uploaded_pdf, folder):
    with open(os.path.join(folder, uploaded_pdf.name), "wb") as f:
        f.write(uploaded_pdf.getbuffer())
    return os.path.join(folder, uploaded_pdf.name)

uploaded_ref_pdf_documents = st.file_uploader("Upload reference documents", type=["pdf"],accept_multiple_files=True)
ref_pdf_folder_name = "ref_pdf_files"
os.makedirs(ref_pdf_folder_name, exist_ok=True)
if uploaded_ref_pdf_documents:
    reference_documents_path = []
    for uploaded_file in uploaded_ref_pdf_documents:
        reference_documents_path.append(save_uploaded_pdf(uploaded_file, ref_pdf_folder_name))



def create_paragraph_from_csv(csv_path):
    paragraph_list = []
    with open('./csv/'+csv_path, errors='ignore') as csvfile:
        rows = csv.DictReader(csvfile)
        for row in rows:
            if len(row["text"]) > 2:
                paragraph_list.append(row["text"].replace('\n', ''))
    return paragraph_list



def extract_text_from_pdf(pdf_path):    # this function will simply load pdf file and return the content inside the file.
        text = ""
        with open(pdf_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
            return text
        


def link_maindoc_references(main_document_path):   # create the link between the citation number and its corresponding reference documents mentioned in the main document
        pattern = r"References([\s\S]*?)(?=(?:\d+\s*\[)|(?:\Z))"
        citation_dict = {}
        references = ""
        match = re.search(pattern, extract_text_from_pdf(main_document_path))                  # Search for the "References" section and extract the content
        if match:
            references += match.group(0)
        lines = references.split('\n')
        for line in lines:
            if line.strip():                          # Check if the item is not an empty string or just whitespace
                parts = line.split('] ', 1)           # Split the item at the first ']' followed by a space
                if len(parts) == 2:
                    nums = parts[0].split('[',1)
                    for num in nums: 
                        if num:
                            key = int(num)          # Extract the reference number (e.g., '[1]')
                            value = parts[1].strip()          # Extract the corresponding text
                            citation_dict[key] = value
            # match = re.match(r'\[(\d+)\](.*)', line)
            # current_citation = int(match.group(1))
            # if match:
            #     # current_citation = int(match.group(1))
            #     citation_dict[current_citation] = match.group(2).strip()
            # else:
            #     if current_citation is not None:
            #         citation_dict[current_citation] += f" {line.strip()}"
        
        return citation_dict



def create_processed_csv():
    csv_path = create_csv("flowchart_to_text_output.pdf")    
    paragraphs = create_paragraph_from_csv(csv_path)
    with open("./csv/output.csv", 'w', newline='') as csvfile:
        fieldnames=["id","text"]
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for id, key in enumerate(paragraphs):
            writer.writerow({"id":id, "text":key})
    csv_path = "output.csv"
    return csv_path



def create_sentences_from_pdf(pdf_path):   # We can't extract each paragraph from the documents. So this function will return a paragraph by combining the continuous rows in the csv.
    csv_path = create_csv(pdf_path)
    paragraphs = create_paragraph_from_csv(csv_path)
    sentences_list = []
    for para in paragraphs:
        sentence_list = sent_tokenize(para)
        for sentence in sentence_list:
            sentences_list.append(sentence)
    
    return sentences_list



citation_pattern1 = r'\[[\d\s,]+\]'
citation_pattern2 = r'\[(\d+(?:,\s*\d+)*)\]'
author_pattern1   = r'\[(?:[A-Z][a-z]*\s)+et al\.,\s*\d{4}\]'  #r'\[(?:[A-Z][a-z]*\s)+et al\.,\s\d{4}\]' 
author_pattern2   = r'\[[A-Za-z\s.&]+ \d{4}\]'          #r'(?:[A-Z][a-z]*\s)+et al\.,\s\d{4}'
table_pattern     = r'Table\s*\d+'
figure_pattern    = r'fig\s*\d+'
http_pattern      = r'https://\S+'



def create_paragraph_reference_dictionary(main_document_path):    # detect all the referencing contents(like citation number[1], authour name, table number, figure number, http link) in the main document and return a dictionary with each paragraph as key and its corresponding referencing content as values)
        csv_path = create_processed_csv()
        paragraphs = create_paragraph_from_csv(csv_path)
        pattern1 = f'({citation_pattern1}|{author_pattern1}|{figure_pattern}|{table_pattern}|{http_pattern})'
        pattern2 = f'({citation_pattern2}|{author_pattern1}|{figure_pattern}|{table_pattern}|{http_pattern})'
        paragraph_reference_dict = {}
        sentences_list = []
        for paragraph in paragraphs:
            sentences = ""
            sentences = re.split(pattern1, paragraph)
            for i in range(0, len(sentences), 2):
                if i + 1 < len(sentences) :
                    sentences_list.append(sentences[i] + sentences[i + 1])
                else:
                    sentences_list.append(sentences[i])
            
            citation_matches = re.findall(pattern2, paragraph)
            if citation_matches:
                for sentence in sentences_list:
                    match = re.search(pattern2, sentence)
                    if match:
                        if len(sentence) > 10:
                            sentence_text = sentence[:match.start()].strip()
                            citations = [c for c in match.group(1).split(',')]
                            paragraph_reference_dict[sentence_text] = ''.join(citations)
                        else:
                            paragraph_reference_dict[sentence] = "[0]"
                    else:
                        paragraph_reference_dict[sentence] = "[0]"
            else:
                paragraph_reference_dict[paragraph] = "[0]"
        return paragraph_reference_dict



def find_similar_document(reference,reference_documents_path):      # identify the correct reference document related to the reference by performing similarity search
    document_texts = [extract_text_from_pdf(pdf) for pdf in reference_documents_path]       # Extract text from PDF documents
    tfidf_vectorizer = TfidfVectorizer()                                                    # Vectorize the text using TF-IDF
    tfidf_matrix = tfidf_vectorizer.fit_transform([reference] + document_texts)
    cosine_similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()  # Calculate cosine similarity between the query text and each document
    most_related_document_index = cosine_similarities.argmax()                              # Find the index of the most related document
    most_related_document_path = reference_documents_path[most_related_document_index]      # Get the path of the most related document
    return most_related_document_path



def map_citation_with_reference_document(cn, reference_documents_path):     # if the referenceing content is citation number then this function will find the correct reference document for citation number reference and create a dictionary with citation number as key and corresponding paragraphs in the referrence document as values
    citation_network = link_maindoc_references(main_document_path)
    reference_dict = {}
    corresponding_document_path = find_similar_document(citation_network[cn],reference_documents_path)
    para_list = create_sentences_from_pdf(corresponding_document_path)
    reference_dict[cn] = para_list
    return reference_dict



def map_author_with_reference_document(reference,reference_documents_path):   #if the referencing content is authors name then this function will find the correct reference document to that authors name and create a dictionary with key as authour name and its corresponding reference documents paragraph as values.
    reference_dict = {}
    corresponding_document_path = find_similar_document(reference,reference_documents_path)
    para_list = create_sentences_from_pdf(corresponding_document_path)
    reference_dict[reference] = para_list
    return reference_dict  



def map_http_link(url):
    response = requests.get(url)
    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'lxml')
        paragraphs = soup.find_all('p')
        page_content = '\n'.join([p.get_text() for p in paragraphs])
        # return page_content
    else:
        print(f"Failed to retrieve content from {url}. Status code: {response.status_code}")
        return None
    return page_content


if "embedding" not in st.session_state:
  st.session_state.embedding = HuggingFaceEmbeddings()


id_db_dict = {}
db_id_dict = {}
if uploaded_ref_pdf_documents:
    for i,j in enumerate(reference_documents_path,start=1):
      sent_list = create_sentences_from_pdf(j)
      db = FAISS.from_texts(sent_list, st.session_state.embedding)   
      id_db_dict[f"db{i}"] = db                                  
      db_id_dict[sent_list[0]] = f"db{i}"  


def similarity_search_from_pdf(paragraph,reference_sentences_list):  #takes paragraph and its corresponding reference document as input and perform sentence-wise similarity search 
    if reference_sentences_list[0] in db_id_dict:
        d_b = id_db_dict[db_id_dict[reference_sentences_list[0]]]
    sentences = sent_tokenize(paragraph)
    lst=[]
    for sentence in sentences:
        similar_docs = d_b.similarity_search(sentence, k=1)
        lst.append(similar_docs[0].page_content)
    return list(set(lst))

def similarity_search_from_paragraph(paragraph,relavant_content):
    relavant_content_lst = sent_tokenize(relavant_content)
    db = FAISS.from_texts(relavant_content_lst, st.session_state.embedding) 
    lst = db.similarity_search(paragraph, k=1)
    lst=[]
    paragraphs = sent_tokenize(paragraph)
    for para in paragraphs:
        similar_docs = db.similarity_search(para, k=1)
        lst.append(similar_docs[0].page_content)
    return list(set(lst))


def create_knowledge_content(main_document_path, reference_documents_path):   # this function will create a final list which contain each paragraph from the main document and relevant contents for that paragraph from the corresponding reference document, only if that paragraph is cited or refered.
    extracted_content = [] 
    paragraph_citations_dictionary = create_paragraph_reference_dictionary(main_document_path)

    for para, ref in paragraph_citations_dictionary.items(): 

        if para in st.session_state.table_id_dict:
            extracted_content.append(st.session_state.table_id_dict[para])

        elif para in st.session_state.flowchart_id_dict:
            extracted_content.append(st.session_state.flowchart_id_dict[para])
        
        elif ref in st.session_state.table_id_dict:
            extracted_content.append(para.upper())
            extracted_content.append(st.session_state.table_id_dict[ref])

        elif ref in st.session_state.flowchart_id_dict:
            extracted_content.append(para.upper())
            extracted_content.append(st.session_state.flowchart_id_dict[ref])

        elif re.findall(citation_pattern1,ref) and ref != "[0]":
            extracted_content.append(para.upper())
            cit = re.findall(citation_pattern1,ref)
            cit = [int(c.strip("[]")) for c in cit]
            relavant_content = map_citation_with_reference_document(cit[0], reference_documents_path)
            result = similarity_search_from_pdf(para,relavant_content[cit[0]])
            extracted_content.append(result)
         
        elif re.findall(author_pattern2,ref):
            extracted_content.append(para.upper())
            cit = re.findall(author_pattern2,ref)
            relavant_content = map_author_with_reference_document(cit[0], reference_documents_path)
            result = similarity_search_from_pdf(para,relavant_content[cit[0]])
            extracted_content.append(result)
         
        elif re.findall(table_pattern,ref):
            extracted_content.append(para.upper())
            cit = re.findall(table_pattern,ref)
            result = st.session_state.table_id_dict[re.sub(r'\s+', ' ', ref)]
            extracted_content.append(result)
         
        elif re.findall(figure_pattern,ref):
            extracted_content.append(para.upper())
            cit = re.findall(figure_pattern,ref)
            result = st.session_state.flowchart_id_dict[re.sub(r'\s+', ' ', ref)]
            extracted_content.append(result)
         
        elif re.findall(http_pattern,ref):
            extracted_content.append(para.upper())
            relavant_content = map_http_link(ref)
            if relavant_content:
                result = similarity_search_from_paragraph(para,relavant_content)
                extracted_content.append(result)
         
        elif ref == "[0]":
            extracted_content.append(para.upper())

        else:
            extracted_content.append("No Reference found")
    # st.write(extracted_content)
    return extracted_content


def create_empty_doc():
    doc = Document()
    return doc

def show_pdf(file_path):
    with open(file_path,"rb") as f:
          base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="900" type="application/pdf">'
    st.markdown(pdf_display, unsafe_allow_html=True)

def append_content_to_docx(doc, content):
    if isinstance(content, str):
        doc.add_paragraph(content, style='BodyText')
    elif isinstance(content, int):
        doc.add_paragraph(str(content), style='BodyText')
    elif isinstance(content, bytes):
        print("entered bytes ===> ", content)
        decoded_str = content.decode('utf-8')
        doc.add_paragraph(decoded_str, style='BodyText')
    elif isinstance(content, tuple):
        print("entered tuple ===> ", content)
        formatted_str = ', '.join(map(str, content))
        doc.add_paragraph(formatted_str, style='BodyText')
    else:
        print("escaped")
        print(content)



if uploaded_main_pdf_document and uploaded_ref_pdf_documents:
    if st.button("Generate"):
        if "table_id_dict" and "flowchart_id_dict" not in st.session_state:
            st.session_state.table_id_dict = table_to_text(main_document_path)  # take "test_doc1.pdf" as input and give "table_to_text.pdf" as output
            
            st.session_state.flowchart_id_dict = flowchart_to_text("table_to_text_output.pdf")  # take "table_to_text_output.pdf" as input and give "final.pdf" as output
            
        result = create_knowledge_content(main_document_path, reference_documents_path)
        empty_doc = create_empty_doc()
        for item in result:
            if type(item) == list:
                for i in item:
                    append_content_to_docx(empty_doc, i)
                empty_doc.save("Knowledge_document.docx")
            else: 
                append_content_to_docx(empty_doc, item)
                empty_doc.save("Knowledge_document.docx")
        convert("Knowledge_document.docx")
        st.success("Knowledge Network created sucessfully")
        show_pdf("Knowledge_document.pdf")
